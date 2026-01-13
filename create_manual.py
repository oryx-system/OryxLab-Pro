"""
지혜마루 도서관 예약 시스템 - 모바일 사용 설명서 생성 스크립트
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Screenshot paths
SCREENSHOTS = {
    'home': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_main_page_calander_1768289592554.png',
    'my_reservation': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_my_reservation_auth_1768289637941.png',
    'checkin': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_checkin_page_1768289735004.png',
    'login': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_login_page_1768289781479.png',
}

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('지혜마루 작은 도서관', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('모바일 사용 설명서')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # ========== Section 1: 예약하기 ==========
    doc.add_heading('1. 예약하기', level=1)
    
    doc.add_paragraph('스마트폰으로 도서관 홈페이지에 접속하면 예약 캘린더가 표시됩니다.')
    
    # Add home screenshot
    if os.path.exists(SCREENSHOTS['home']):
        doc.add_picture(SCREENSHOTS['home'], width=Cm(8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    steps = doc.add_paragraph()
    steps.add_run('📱 예약 방법:\n').bold = True
    steps.add_run('① 원하는 날짜를 터치하세요.\n')
    steps.add_run('② 예약 가능한 시간대가 표시됩니다.\n')
    steps.add_run('③ 시간대를 선택하고 정보를 입력하세요.\n')
    steps.add_run('④ "예약하기" 버튼을 터치하면 완료!')
    
    doc.add_page_break()
    
    # ========== Section 2: 내 예약 확인 ==========
    doc.add_heading('2. 내 예약 확인하기', level=1)
    
    doc.add_paragraph('상단 메뉴에서 "내 예약"을 터치하면 예약 내역을 조회할 수 있습니다.')
    
    # Add my_reservation screenshot
    if os.path.exists(SCREENSHOTS['my_reservation']):
        doc.add_picture(SCREENSHOTS['my_reservation'], width=Cm(8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    steps = doc.add_paragraph()
    steps.add_run('📱 조회 방법:\n').bold = True
    steps.add_run('① 예약 시 입력한 전화번호를 입력하세요.\n')
    steps.add_run('② 4자리 비밀번호를 입력하세요.\n')
    steps.add_run('③ "조회하기" 버튼을 터치하세요.\n')
    steps.add_run('④ 예약 내역과 상세 정보가 표시됩니다.')
    
    doc.add_page_break()
    
    # ========== Section 3: QR 체크인 ==========
    doc.add_heading('3. QR 체크인 (입실 확인)', level=1)
    
    doc.add_paragraph('도서관 도착 후, 출입문에 부착된 QR 포스터를 스캔하여 입실을 확인합니다.')
    
    # Add checkin screenshot
    if os.path.exists(SCREENSHOTS['checkin']):
        doc.add_picture(SCREENSHOTS['checkin'], width=Cm(8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    steps = doc.add_paragraph()
    steps.add_run('📱 체크인 방법:\n').bold = True
    steps.add_run('① 스마트폰의 기본 카메라 앱을 열어주세요.\n')
    steps.add_run('② 출입문의 QR 포스터를 비춰주세요.\n')
    steps.add_run('③ 자동으로 체크인 페이지가 열립니다.\n')
    steps.add_run('④ 전화번호와 비밀번호를 확인하세요.\n')
    steps.add_run('⑤ "입실 확인" 버튼을 터치하면 완료!')
    
    doc.add_paragraph()
    
    note = doc.add_paragraph()
    note.add_run('💡 참고: ').bold = True
    note.add_run('체크인은 예약 시간 30분 전부터 당일 자정까지 가능합니다.')
    
    doc.add_page_break()
    
    # ========== Section 4: 관리자 로그인 (선택) ==========
    doc.add_heading('4. 관리자 페이지 (관리자용)', level=1)
    
    doc.add_paragraph('관리자는 별도의 비밀번호로 로그인하여 예약을 관리할 수 있습니다.')
    
    # Add login screenshot
    if os.path.exists(SCREENSHOTS['login']):
        doc.add_picture(SCREENSHOTS['login'], width=Cm(8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    steps = doc.add_paragraph()
    steps.add_run('🔐 관리자 기능:\n').bold = True
    steps.add_run('• 전체 예약 목록 조회 및 검색\n')
    steps.add_run('• 예약 취소 및 상태 변경\n')
    steps.add_run('• 노쇼 사용자 차단 관리\n')
    steps.add_run('• 통계 및 이용 현황 확인\n')
    steps.add_run('• 공지사항, 비밀번호, QR 코드 설정')
    
    doc.add_page_break()
    
    # ========== Section 5: 문의 및 주의사항 ==========
    doc.add_heading('5. 이용 안내 및 주의사항', level=1)
    
    points = [
        ('⏰ 이용 시간', '예약한 시간 내에만 이용 가능합니다.'),
        ('📱 체크인 필수', '입실 시 반드시 QR 체크인을 해주세요.'),
        ('🚫 노쇼 주의', '예약 후 미이용 시 이용이 제한될 수 있습니다.'),
        ('📸 퇴실 인증', '퇴실 시 정리 사진을 촬영해주세요.'),
        ('🔒 개인정보', '입력하신 정보는 예약 관리 목적으로만 사용됩니다.'),
    ]
    
    for title, desc in points:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph('© 2026 지혜마루 작은 도서관')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].font.italic = True
    
    powered = doc.add_paragraph('Powered by oryxdatasafe')
    powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    powered.runs[0].font.size = Pt(9)
    powered.runs[0].font.italic = True
    
    # Save
    output_path = r'C:\Workspace\OryxLab_Pro\지혜마루_모바일_사용설명서.docx'
    doc.save(output_path)
    print(f'✅ 사용 설명서가 생성되었습니다: {output_path}')
    return output_path

if __name__ == '__main__':
    create_manual()
