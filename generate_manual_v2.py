from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_detailed_manual():
    doc = Document()
    
    # 1. Styles Setup (Korean Font)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Helper: Add Title
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER # Spacer

    # Helper: Add Chapter Header
    def add_chapter(title):
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204) # Blue
        p.space_after = Pt(12)

    # Helper: Add Step
    def add_step(step_num, text, image_file=None, emphasis=None):
        p = doc.add_paragraph()
        
        # Step Number
        step_run = p.add_run(f"Step {step_num}. ")
        step_run.bold = True
        step_run.font.color.rgb = RGBColor(255, 102, 0) # Orange
        
        # Main Text
        p.add_run(text)
        
        # Emphasis text (red)
        if emphasis:
            emp_run = p.add_run(f"\n{emphasis}")
            emp_run.bold = True
            emp_run.font.color.rgb = RGBColor(255, 0, 0)

        # Image
        if image_file:
            img_path = os.path.join(os.getcwd(), 'screenshots_v2', image_file)
            if os.path.exists(img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(img_path, width=Inches(3.5))
                # Add border effect description roughly
                img_p.space_after = Pt(12)
            else:
                 doc.add_paragraph(f"[이미지를 찾을 수 없습니다: {image_file}]")

    # --- Document Content ---

    # Cover Page
    add_title("지혜마루 작은 도서관\n모바일 앱 사용 설명서")
    doc.add_paragraph("누구나 쉽게 따라할 수 있는 단계별 가이드입니다.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("이 설명서는 다음 내용을 포함합니다:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc = [
        "1. 처음 화면 (공지사항 확인)",
        "2. 예약하기 (가장 중요!)",
        "3. 전자 서명하는 법",
        "4. 입장하기 (체크인)",
        "5. 내 이용 내역 확인하기"
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Chapter 1: Main
    add_chapter("1. 처음 화면 살펴보기")
    add_step(1, 
             "앱에 처음 접속하면 아래와 같은 화면이 나옵니다. 도서관의 공지사항을 확인하실 수 있습니다.", 
             "1_Main.png")
    add_step(2, 
             "화면 아래쪽에 있는 메뉴 버튼을 주목해주세요. '체크인' 버튼과 '마이 페이지'버튼이 있습니다.")

    # Chapter 2: Reservation
    add_chapter("2. 예약하기 (따라해보세요)")
    doc.add_paragraph("도서관을 이용하려면 먼저 예약을 해야 합니다. 천천히 따라해보세요.")
    
    add_step(1, 
             "달력에서 원하시는 날짜를 손가락으로 꾹 눌러주세요.", 
             "1_Main.png", 
             "팁: 오늘 날짜나 내일 날짜를 선택해보세요.")
    
    add_step(2, 
             "날짜를 누르면 예약 정보를 입력하는 창이 뜹니다. 성함과 전화번호 뒷자리를 입력해주세요.", 
             "2_Reservation_Modal.png")

    add_step(3, 
             "빈칸을 모두 채워주세요. 비밀번호는 나중에 확인할 때 필요하니 꼭 기억해주세요!",
             "3_Reservation_Filled.png",
             "중요: 전화번호와 비밀번호를 잊어버리면 조회가 어렵습니다.")

    # Chapter 3: Signature
    add_chapter("3. 전자 서명하기 (중요)")
    add_step(1, 
             "예약 정보 아래쪽에 하얀색 네모난 공간이 있습니다. 이곳이 서명하는 곳입니다.", 
             "4_Reservation_Signature.png")
    
    add_step(2, 
             "손가락으로 본인의 이름을 정자로, 크게 적어주세요.", 
             None,
             "주의: 서명을 하지 않으면 예약이 완료되지 않습니다.")
    
    add_step(3, 
             "모두 입력하셨다면 맨 아래 '예약하기' 버튼을 눌러주세요.")

    # Chapter 4: Check-in
    add_chapter("4. 도서관 입장하기 (체크인)")
    add_step(1, 
             "도서관 문 앞에 도착하셨나요? 출입문에 붙어있는 QR 코드를 찾아주세요.")
    
    add_step(2, 
             "앱 화면 아래쪽의 '체크인' 메뉴를 누르신 후, 전화번호와 비밀번호를 입력해주세요.", 
             "5_Checkin_Filled.png")
    
    add_step(3, 
             "'QR 스캔하기' 버튼을 누르고, 카메라로 출입문의 QR 코드를 비춰주세요. 문이 열립니다!")

    # Chapter 5: My Page
    add_chapter("5. 내 예약 확인하기")
    add_step(1, 
             "내가 언제 예약했는지 깜빡하셨나요? '마이 페이지' 메뉴를 눌러보세요.", 
             "6_MyPage.png")
    
    add_step(2, 
             "전화번호로 조회하면 나의 예약 내역과, 도서관 와이파이 비밀번호 등을 확인할 수 있습니다.")


    # Save
    save_path = os.path.join(os.getcwd(), '지혜마루_사용설명서_상세본.docx')
    doc.save(save_path)
    print(f"Detailed Manual generated at: {save_path}")

if __name__ == "__main__":
    create_detailed_manual()
