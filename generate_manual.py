from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_manual():
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading('지혜마루 작은 도서관 이용 가이드', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('모바일에 최적화된 지혜마루 작은 도서관 앱 사용법을 안내합니다.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER # Spacer

    # Helper to add section
    def add_section(title, image_filename, description_list):
        doc.add_heading(title, level=1)
        
        # Add Screenshot
        img_path = os.path.join(os.getcwd(), 'screenshots', image_filename)
        if os.path.exists(img_path):
            # Center the image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(3.0)) # Mobile width optimized
        else:
            doc.add_paragraph(f"[Image not found: {image_filename}]", style='Body Text')

        # Add Description
        for desc in description_list:
            p = doc.add_paragraph(desc, style='List Bullet')

        doc.add_page_break()

    # Section 1: 홈 화면
    add_section(
        "1. 홈 화면 (메인)",
        "1_main.png",
        [
            "앱 접속 시 첫 화면입니다.",
            "공지사항을 확인할 수 있습니다.",
            "하단 메뉴를 통해 '체크인' 및 '마이페이지'로 이동할 수 있습니다."
        ]
    )

    # Section 2: 체크인
    add_section(
        "2. 체크인 (입장하기)",
        "2_checkin.png",
        [
            "도서관 입구의 QR 코드를 스캔하세요.",
            "전화번호와 비밀번호를 입력 후 '입장하기' 버튼을 누르세요.",
            "예약 시간 30분 전부터 체크인이 가능합니다."
        ]
    )

    # Section 3: 마이 페이지
    add_section(
        "3. 마이 페이지 (이용 내역)",
        "3_mypage.png",
        [
            "나의 예약 현황을 확인할 수 있습니다.",
            "예약 취소 및 패널티 현황을 조회할 수 있습니다.",
            "와이파이 정보 및 출입 비밀번호를 확인할 수 있습니다."
        ]
    )
    
     # Section 4: 로그인 (관리자)
    add_section(
        "4. 관리자 / 개발자 로그인",
        "4_login.png",
        [
            "관리자 모드로 진입하기 위한 로그인 화면입니다.",
            "일반 사용자는 사용할 필요가 없습니다."
        ]
    )


    # Auto-saved path
    save_path = os.path.join(os.getcwd(), 'OryxLab_User_Manual.docx')
    doc.save(save_path)
    print(f"User Manual generated at: {save_path}")

if __name__ == "__main__":
    create_manual()
