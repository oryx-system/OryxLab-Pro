"""
지혜마루 도서관 "친절한" 사용 설명서 생성 스크립트
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# Screenshot paths
SCREENSHOTS = {
    'home': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_main_page_calander_1768289592554.png',
    'res_form': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_reservation_modal_1768290108424.png',
    'sig_modal': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\signature_modal_opened_1768291005163.png',
    'sig_done': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\reservation_form_with_signature_preview_1768291035307.png',
    'my_auth': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_my_reservation_auth_1768289637941.png',
    'my_list': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_my_reservation_list_1768291253269.png',
    'checkin_page': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_checkin_page_1768289735004.png',
    'checkin_success': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_checkin_success_1768291280492.png',
    'admin_login': r'C:\Users\hongs\.gemini\antigravity\brain\670da5e4-2e96-4f71-b5ba-260edf85ea43\mobile_login_page_1768289781479.png',
}

def add_heading_style(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0]
    run.font.name = "Malgun Gothic"
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def add_image_centered(doc, path, width_cm=7.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add a little space
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    else:
        doc.add_paragraph(f"[이미지 없음: {os.path.basename(path)}]")

def add_step_text(doc, number, text):
    p = doc.add_paragraph()
    runner = p.add_run(f"Step {number}. ")
    runner.bold = True
    runner.font.size = Pt(13)
    runner.font.color.rgb = RGBColor(0, 102, 204) # Lighter Blue
    
    runner2 = p.add_run(text)
    runner2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def create_manual():
    doc = Document()
    
    # --- Title ---
    title = doc.add_heading('지혜마루 사용자 가이드', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('누구나 쉽게 따라하는 예약/입실 방법', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Intro ---
    doc.add_paragraph("안녕하세요! 지혜마루 도서관입니다.\n스마트폰으로 간편하게 예약을 하고, 출입문에서 QR코드를 찍어 입장하는 방법을 아주 쉽게 알려드릴게요.", style='List Paragraph')
    doc.add_page_break()

    # --- Part 1: 예약하기 ---
    add_heading_style(doc, '1. 자리 예약하기', 1)
    
    add_step_text(doc, 1, "홈페이지에 접속하면 달력이 나옵니다.")
    add_image_centered(doc, SCREENSHOTS['home'])
    doc.add_paragraph("원하는 날짜를 손가락으로 톡! 눌러주세요.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_step_text(doc, 2, "예약 정보를 입력해주세요.")
    add_image_centered(doc, SCREENSHOTS['res_form'])
    doc.add_paragraph("이름과 전화번호를 정확히 적어주세요.\n(비밀번호 4자리는 꼭 기억해주세요!)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    
    # --- Part 2: 전자 서명 (중요) ---
    add_heading_style(doc, '2. 전자 서명하기 (중요 ✍️)', 1)
    doc.add_paragraph("예약 마지막 단계에서 꼭 '서명'을 해야 합니다.")

    add_step_text(doc, 1, "'서명하기' 버튼을 누르세요.")
    add_image_centered(doc, SCREENSHOTS['sig_modal'])
    doc.add_paragraph("화면의 네모 칸 안에 손가락으로 이름을 써주세요.\n다 쓰셨으면 아래 '완료' 버튼을 누릅니다.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_step_text(doc, 2, "서명이 잘 들어갔는지 확인하세요.")
    add_image_centered(doc, SCREENSHOTS['sig_done'])
    doc.add_paragraph("서명이 보이면, 맨 아래 '예약 신청' 버튼을 눌러 예약을 끝냅니다.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- Part 3: 예약 확인 및 QR 체크인 ---
    add_heading_style(doc, '3. 내 예약 확인 & 입실하기', 1)

    add_step_text(doc, 1, "'내 예약' 메뉴에서 전화번호를 입력하세요.")
    add_image_centered(doc, SCREENSHOTS['my_auth'])
    
    add_step_text(doc, 2, "예약 목록이 나타납니다.")
    add_image_centered(doc, SCREENSHOTS['my_list'])
    doc.add_paragraph("예약 날짜와 시간을 확인할 수 있습니다.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    add_heading_style(doc, '4. 도서관 도착! 입실 확인', 1)
    doc.add_paragraph("도서관 문 앞에 붙어있는 QR코드를 찍어야 문이 열립니다!")
    
    add_step_text(doc, 1, "카메라로 QR을 찍으면 체크인 화면이 뜹니다.")
    add_image_centered(doc, SCREENSHOTS['checkin_page'])
    doc.add_paragraph("자동으로 체크인 화면이 열립니다.\n'입실 확인' 버튼을 꾹 눌러주세요.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_step_text(doc, 2, "입실 완료 메세지가 뜨면 성공!")
    add_image_centered(doc, SCREENSHOTS['checkin_success'])
    doc.add_paragraph("이제 도서관을 자유롭게 이용하시면 됩니다. 환영합니다!").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Part 4: 관리자 ---
    add_heading_style(doc, '5. 관리자 로그인 (관계자용)', 1)
    add_image_centered(doc, SCREENSHOTS['admin_login'])
    doc.add_paragraph("관리자 페이지는 승인된 관리자만 접속 가능합니다.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('문의: 지혜마루 도서관 관리실 (010-0000-0000)')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = r'C:\Workspace\OryxLab_Pro\지혜마루_상세_사용설명서.docx'
    doc.save(output_path)
    print(f"Manual Created: {output_path}")

if __name__ == '__main__':
    create_manual()
